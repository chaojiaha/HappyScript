#coding:utf-8
import os
import datetime

from selenium import webdriver
from PIL import Image
from docx import Document
from docx.shared import Inches
import requests


INVOICENO = ""
REVENUEREGISTERID = ""

def take_screenshot(browser, save_fn="capture.png"):
    '''截图'''
    browser.save_screenshot(save_fn)
    im = Image.open(save_fn)
    box=(400,30,1100,600)
    im = im.crop(box)
    im.save(save_fn)

def edit_page(browser, base_string, edit_to_string):
    '''修改静态页面'''
    script = '''
    document.getElementById("jgms").innerHTML = document.getElementById("jgms").innerHTML.replace("{}", "{}")
    '''.format(base_string, edit_to_string)
    browser.execute_script(script)
    return browser

def convert_image(image):
    '''图片处理'''
    image2=image.convert('L')
    return image2

def get_captcha(cookies):
    '''获取验证码'''
    cookie = {"JSESSIONID": cookies["value"]}
    image_url = "http://www.tax.sh.gov.cn/servlet/GetshowimgSmall"
    image_re = requests.get(image_url, cookies=cookie)
    with open('image_base.jpg', 'w') as f:
        f.write(image_re.content)
    im = Image.open('image_base.jpg')
    im = convert_image(im)
    im.show()
    yzm = input("input yzm: ")
    return yzm

def save_to_doc(img_name):
    document = Document()
    document.add_picture(img_name)
    return document

def main():
    start = str(input("start from: "))
    end = str(input("end to:"))
    # base_url = "http://www.tax.sh.gov.cn/wsbs/WSBSptFpCx_loginsNewl.jsp"
    base_url = "https://www.tax.sh.gov.cn/xbwz/wzcx/WSBSptFpCx_loginsNewl.jsp"
    browser = webdriver.Chrome()
    browser.maximize_window()
    browser.get(base_url)
    # browser.refresh()
    try:
        # 下三行把验证码提取出来可以做响应处理，验证码识别
        # cookies = browser.get_cookie("JSESSIONID")
        # print cookies
        # captcha = get_captcha(cookies)
        captcha = input("input yzm: ")
        browser.find_element_by_id("invoiceNo").send_keys(INVOICENO)
        browser.find_element_by_id("fphm").send_keys(start)
        browser.find_element_by_id("revenueRegisterId").send_keys(REVENUEREGISTERID)
        browser.find_element_by_id("yzm").send_keys(captcha)
        browser.find_element_by_name("confirm").click()

        document = Document()

        base_string = start
        count = 0
        for i in range(int(start), int(end) + 1):
            edit_to_string = "0"*(8-len(str(i))) + str(i) if len(str(i)) <= 7 else str(i)
            edit_page(browser, base_string, edit_to_string)
            take_screenshot(browser, edit_to_string +'.png')
            document.add_picture(edit_to_string +'.png', width=Inches(4.25))
            os.remove(edit_to_string +'.png')
            base_string = edit_to_string
            count += 1
            print(count, "/", int(end) + 1 -int(start))

        docName = 'taxVerify-' + datetime.datetime.now().strftime('%Y-%m-%d-%H-%M-%S') + '.docx'
        document.save(docName)

    except Exception:
        print("Error!")

    finally:
        browser.close()

if __name__ == '__main__':
    main()